from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_report(data_summary: dict, charts: list[bytes], tables: list[list], config: dict) -> bytes:
    """生成 Word 报告"""
    doc = Document()

    # 标题
    title = doc.add_heading(config.get('title', '数据分析报告'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据概要
    doc.add_heading('数据概要', level=1)
    for key, value in data_summary.items():
        doc.add_paragraph(f'{key}: {value}')

    # 图表
    doc.add_heading('分析图表', level=1)
    for chart_bytes in charts:
        doc.add_picture(io.BytesIO(chart_bytes), width=Inches(5.5))

    # 表格
    doc.add_heading('数据表格', level=1)
    for table_data in tables:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)

    # 保存
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()